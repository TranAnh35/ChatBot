import os
from fastapi import UploadFile
from docling.document_converter import DocumentConverter
from docx import Document
import yaml
from typing import List
import io

UPLOAD_FOLDER = "upload"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def save_file(file: UploadFile) -> str:
    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    with open(file_path, "wb") as f:
        f.write(file.file.read())
    return file_path

def delete_file(file_name: str) -> bool:
    file_path = os.path.join(UPLOAD_FOLDER, file_name)
    if os.path.exists(file_path):
        os.remove(file_path)
        return True
    return False

def list_files() -> List[dict]:
    files = []
    for file_name in os.listdir(UPLOAD_FOLDER):
        file_path = os.path.join(UPLOAD_FOLDER, file_name)
        size = os.path.getsize(file_path)
        files.append({"name": file_name, "size": size, "path": file_path})
    return files

"""
* Các file đọc nội dung từ file được lưu trong thư mục upload (sử dụng cho RAG)
"""

def read_file_from_path(file_path: str) -> str:
    """Đọc nội dung file dựa trên định dạng."""
    file_name = os.path.basename(file_path)
    if file_name.endswith('.pdf'):
        return read_pdf_from_path(file_path)
    elif file_name.endswith(('.doc', '.docx')):
        return read_docx_from_path(file_path)
    elif file_name.endswith(('.yaml', '.yml')):
        return read_yaml_from_path(file_path)
    elif file_name.endswith(('.txt', '.md')):
        return read_txt_from_path(file_path)
    else:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

def read_pdf_from_path(file_path: str) -> str:
    """Đọc nội dung từ file PDF với khả năng xử lý PDF phức tạp sử dụng docling."""
    try:
        # Sử dụng docling DocumentConverter để trích xuất văn bản
        converter = DocumentConverter()
        result = converter.convert(file_path)
        return result.document.export_to_markdown()
    except Exception as e:
        print(f"Lỗi khi sử dụng docling: {str(e)}.")
        raise Exception(f"Không thể đọc file PDF: {str(e)}")

def read_docx_from_path(file_path: str) -> str:
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def read_yaml_from_path(file_path: str) -> str:
    with open(file_path, 'r', encoding='utf-8') as f:
        data = yaml.safe_load(f)
        return str(data)

def read_txt_from_path(file_path: str) -> str:
    encodings = ['utf-8', 'latin1', 'cp1252', 'iso-8859-1']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise Exception(f"Không thể đọc file với các encoding đã thử")

async def read_uploaded_file(file: UploadFile) -> str:
    """Đọc nội dung file được gửi trực tiếp từ frontend dựa trên định dạng."""
    file_name = file.filename
    content: str = ""
    
    try:
        if file_name.endswith('.pdf'):
            content = await read_pdf_from_upload(file)
        elif file_name.endswith(('.doc', '.docx')):
            content = await read_docx_from_upload(file)
        elif file_name.endswith(('.yaml', '.yml')):
            content = await read_yaml_from_upload(file)
        else:
            content = await read_txt_from_upload(file)
        return content
    except Exception as e:
        raise Exception(f"Không thể đọc file {file_name}: {str(e)}")
    finally:
        await file.close()  # Đóng file sau khi đọc

async def read_pdf_from_upload(file: UploadFile) -> str:
    """Đọc nội dung file PDF từ UploadFile với docling."""
    try:
        # Đọc nội dung file
        file_content = await file.read()
        
        # Lưu tạm vào file tạm thời để docling xử lý
        temp_file_path = os.path.join(UPLOAD_FOLDER, "temp_" + file.filename)
        with open(temp_file_path, "wb") as temp_file:
            temp_file.write(file_content)
        
        # Sử dụng docling
        converter = DocumentConverter()
        result = converter.convert(temp_file_path)
        text = result.document.export_to_markdown()
        
        # Xóa file tạm
        os.remove(temp_file_path)
        
        return text
    except Exception as e:
        print(f"Lỗi khi sử dụng docling: {str(e)}.")
        raise Exception(f"Không thể đọc file PDF: {str(e)}")

async def read_docx_from_upload(file: UploadFile) -> str:
    """Đọc nội dung file DOCX từ UploadFile."""
    doc = Document(file.file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

async def read_yaml_from_upload(file: UploadFile) -> str:
    """Đọc nội dung file YAML từ UploadFile."""
    content = await file.read()
    data = yaml.safe_load(content)
    return str(data)

async def read_txt_from_upload(file: UploadFile) -> str:
    """Đọc nội dung file text từ UploadFile."""
    content = await file.read()
    encodings = ['utf-8', 'latin1', 'cp1252', 'iso-8859-1']
    for encoding in encodings:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise Exception(f"Không thể đọc file với các encoding đã thử")